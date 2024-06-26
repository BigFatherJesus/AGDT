import os
import time
import io
import openai
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
from docx import Document
from PyPDF2 import PdfFileReader, PdfFileWriter
from reportlab.pdfgen import canvas
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()  # Load environment variables from .env file

os.environ['OPENAI_API_KEY'] = os.getenv('OPENAI_API_KEY')
client = OpenAI()

# Setup Google Drive API
SCOPES = ['https://www.googleapis.com/auth/drive']
SERVICE_ACCOUNT_FILE = './credentials/credentials.json'

credentials = service_account.Credentials.from_service_account_file(
    SERVICE_ACCOUNT_FILE, scopes=SCOPES)

service = build('drive', 'v3', credentials=credentials)

# Function to list files in a folder
def list_files(folder_id, suffix=None):
    """List all files in the specified Google Drive folder."""
    query = f"'{folder_id}' in parents and trashed=false"
    if suffix:
        query += f" and name ends with '{suffix}'"
    results = service.files().list(
        q=query,
        pageSize=None, fields="files(id, name, mimeType)").execute()
    items = results.get('files', [])
    return items

# Function for downloading files from the drive
def download_file(file_id, file_name):
    """Download a file from Google Drive."""
    request = service.files().get_media(fileId=file_id)
    data_folder = os.path.join(os.getcwd(), "Data")
    if not os.path.exists(data_folder):
        os.makedirs(data_folder)
    file_path = os.path.join(data_folder, file_name)
    with open(file_path, 'wb') as fh:
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
            print(f"Downloading: {int(status.progress() * 100)}%")

# Function to upload the translated files
def upload_file(file_path, folder_id):
    file_name = os.path.basename(file_path)
    file_metadata = {'name': file_name, 'parents': [folder_id]}
    media = MediaFileUpload(file_path, resumable=True)
    file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
    print(f'File ID: {file.get("id")}')

# Function for translating the file --W.I.P--
# We need to cut the file every 125 lines to stay within the token limit
def translate_text(text):
    lines = text.split("\n")
    chunks = []
    chunk = []
    for i, line in enumerate(lines):
        chunk.append(line)
        if (i + 1) % 125 == 0 or i == len(lines) - 1:
            chunks.append("\n".join(chunk))
            chunk = []

    translated_chunks = []
    for chunk in chunks:
        response = client.chat.completions.create(
            model="gpt-4o",
            # System prompt for the AI to understand it's purpose
            messages=[
                {"role": "system", "content": os.getenv('SYSTEM_PROMPT')},
                {"role": "user", "content": chunk}
            ]
        )
        # Remove the "plaintext" text formatting from the AI's output
        choices = response.choices
        translated_chunk = choices[0].message.content
        translated_chunk = translated_chunk.replace("```plaintext", "").replace("```", "").replace("```", "")
        translated_chunks.append(translated_chunk)

    translated_text = "\n".join(translated_chunks)
    return translated_text

# Function for extracting and repackaging the text
def translate_file(file_name, file_type):
    data_folder = os.path.join(os.getcwd(), "Data")
    file_path = os.path.join(data_folder, file_name)
    
    file_path, file_name_only = os.path.split(file_path)
    file_name_without_extension, file_extension = os.path.splitext(file_name_only)
    
    # Support for multiple file types [NOT TESTED, SO FAR ONLY .SRT & .TXT ARE TESTED]
    # Append _AT_Translated at the end of the file name to clarify which are transalted

    # Function for word documents because you never know
    if file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':  # .docx
        doc = Document(file_path)
        translated_paragraphs = []
        for paragraph in doc.paragraphs:
            text = paragraph.text
            translated_text = translate_text(text)
            translated_paragraphs.append(translated_text)

        new_file_name = f"{file_name_without_extension}_AT_Translated{file_extension}"
        new_file_path = os.path.join(file_path, new_file_name)
        new_doc = Document()
        for paragraph in translated_paragraphs:
            new_doc.add_paragraph(paragraph)

        new_doc.save(new_file_path)

    # Function for PDF
    elif file_type == 'application/pdf':  # .pdf
        pdf_reader = PdfFileReader(open(file_path, 'rb'))
        new_file_name = f"{file_name_without_extension}_AT_Translated{file_extension}"
        new_file_path = os.path.join(file_path, new_file_name)
        new_pdf = PdfFileWriter()

        for page_num in range(pdf_reader.numPages):
            page = pdf_reader.getPage(page_num)
            text = page.extractText()
            translated_text = translate_text(text)

            new_page = canvas.Canvas(new_file_path, pagesize=page.mediaBox.getWidth(), bottomup=0)
            new_page.setFont("Helvetica", 12)
            new_page.drawString(100, 700, translated_text)
            new_page.showPage()
            new_page.save()

        with open(new_file_path, 'wb') as out_pdf:
            new_pdf.write(out_pdf)
    # Function for text based files
    else:  # .txt, .srt, etc.
        with open(os.path.join(file_path, file_name_only), 'r', encoding='utf-8') as file:
            text = file.read()
            translated_text = translate_text(text)

        new_file_name = f"{file_name_without_extension}_AT_Translated{file_extension}"
        new_file_path = os.path.join(file_path, new_file_name)
        with open(new_file_path, 'w', encoding='utf-8') as file:
            file.write(translated_text)

    return new_file_path

# Function to list folder contents
def list_folder_contents(folder_id):
    """List the contents of a Google Drive folder."""
    results = service.files().list(
        q=f"'{folder_id}' in parents and trashed=false",
        pageSize=10, fields="files(id, name, mimeType)").execute()
    items = results.get('files', [])
    print(f"Contents of folder {folder_id}:")
    for item in items:
        print(f"  {item['name']} ({item['mimeType']})")

# .env variables, there are more of these, including in the system prompt to make it easier adjustable
input_folder_id = os.getenv('INPUT_FOLDER_ID')
output_folder_id = os.getenv('OUTPUT_FOLDER_ID')

# Function to calculate the amount of files that still need translating
def files_needing_translation(total_files, skipped_files):
    return total_files - skipped_files

# Main function
def main():
    print("Script started...")

    print("Input folder contents:")
    list_folder_contents(input_folder_id)
    print()
    print("Output folder contents:")
    list_folder_contents(output_folder_id)
    print()

    processed_files = set()
    total_files = 0
    skipped_files = 0

    while True:
        files = list_files(input_folder_id)
        total_files = len(files)
        for file in files:
            file_id = file['id']
            file_name = file['name']
            mime_type = file['mimeType']

            # Check if the file already has the "_AT_Translated" suffix, if this is not done it will indefinitly translate one and the same file if on accident the output file is put back into the input, ask me how I know
            if file_name.endswith("_AT_Translated"):
                skipped_files += 1
                continue

            # Check if a file with the same name (without the suffix) already exists in the output folder
            query = f"'{output_folder_id}' in parents and trashed=false and name contains '_AT_Translated'"
            existing_files = service.files().list(q=query, fields="files(id, name, mimeType)").execute().get('files', [])
            existing_file = next((f for f in existing_files if f['name'].replace("_AT_Translated", "") == file_name), None)
            if existing_file:
                skipped_files += 1
                continue

            if file_id not in processed_files:
                files_to_translate = files_needing_translation(total_files, skipped_files)
                print(f"{files_to_translate} file(s) needs translating")
                print(f"Processing file {file_name}...")
                download_file(file_id, file_name)
                print(f"Translating {file_name}...")
                new_file_path = translate_file(file_name, mime_type)
                print(f"Uploading {file_name}...")
                upload_file(new_file_path, output_folder_id)
                print(f"File {file_name} succesfully processed.")
                processed_files.add(file_id)

        print(f"")
        print(f"Skipping {skipped_files}/{total_files} translated files.")
        print(f"Waiting...")
        print(f"")
        time.sleep(60)  # Wait for 1 minute before checking for new files again, this could be changed to a lower / higher number but I have found 1 min to work best
        skipped_files = 0  # Reset the skipped_files counter

# Main loop
if __name__ == '__main__':
    main()